"""
Word文档处理器
支持读取Word文档和添加批注
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import win32com.client as win32
import pythoncom


class WordProcessor:
    """Word文档处理器"""

    def __init__(self, file_path: str = None):
        self.file_path = file_path
        self.document: Optional[Document] = None
        self._word_app = None
        self._word_doc = None
        # 页码估算参数
        self._chars_per_page = 1500  # 每页约1500字符（中文字符）

    def load_document(self, file_path: str = None):
        """加载Word文档"""
        if file_path:
            self.file_path = file_path

        if not self.file_path:
            raise ValueError("文件路径不能为空")

        if not Path(self.file_path).exists():
            raise FileNotFoundError(f"文件不存在: {self.file_path}")

        self.document = Document(self.file_path)

    def _estimate_page_number(self, paragraph_index: int) -> int:
        """估算段落在哪一页"""
        if not self.document:
            return 1

        # 计算到目标段落的总字符数
        total_chars = 0
        for i, para in enumerate(self.document.paragraphs):
            if i >= paragraph_index:
                break
            total_chars += len(para.text)

        # 估算页码
        page = max(1, (total_chars // self._chars_per_page) + 1)
        return page

    def get_paragraphs(self) -> List[Dict[str, Any]]:
        """获取所有段落"""
        if not self.document:
            raise ValueError("文档未加载")

        paragraphs = []
        for i, para in enumerate(self.document.paragraphs):
            paragraphs.append({
                'index': i,
                'text': para.text,
                'style': para.style.name if para.style else None,
                'alignment': para.alignment,
                'runs': [{'text': run.text, 'font': run.font.name} for run in para.runs],
                'estimated_page': self._estimate_page_number(i)
            })

        return paragraphs
    
    def get_tables(self) -> List[Dict[str, Any]]:
        """获取所有表格"""
        if not self.document:
            raise ValueError("文档未加载")
        
        tables = []
        for i, table in enumerate(self.document.tables):
            table_data = {
                'index': i,
                'rows': [],
                'caption': None
            }
            
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data['rows'].append(row_data)
            
            tables.append(table_data)
        
        return tables
    
    def get_headings(self) -> List[Dict[str, Any]]:
        """获取所有标题"""
        if not self.document:
            raise ValueError("文档未加载")

        headings = []
        for i, para in enumerate(self.document.paragraphs):
            if para.style and 'Heading' in para.style.name:
                level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 1
                headings.append({
                    'index': i,
                    'text': para.text,
                    'level': level,
                    'style': para.style.name,
                    'estimated_page': self._estimate_page_number(i)
                })

        return headings
    
    def get_images(self) -> List[Dict[str, Any]]:
        """获取所有图片"""
        if not self.document:
            raise ValueError("文档未加载")
        
        images = []
        # python-docx 无法直接获取图片，需要通过 COM 接口
        # 这里返回一个占位列表
        return images
    
    def get_document_info(self) -> Dict[str, Any]:
        """获取文档信息"""
        if not self.document:
            raise ValueError("文档未加载")
        
        core_props = self.document.core_properties
        
        return {
            'title': core_props.title or "未命名文档",
            'author': core_props.author or "未知作者",
            'subject': core_props.subject,
            'keywords': core_props.keywords,
            'created': core_props.created,
            'modified': core_props.modified,
            'paragraph_count': len(self.document.paragraphs),
            'table_count': len(self.document.tables),
            'section_count': len(self.document.sections)
        }
    
    def add_comment_win32(self, paragraph_index: int, comment_text: str, author: str = "WordChecker"):
        """
        使用Win32 COM接口添加批注（仅Windows）
        
        Args:
            paragraph_index: 段落索引
            comment_text: 批注内容
            author: 批注作者
        """
        try:
            pythoncom.CoInitialize()
            
            # 启动Word应用
            if not self._word_app:
                self._word_app = win32.Dispatch("Word.Application")
            
            # 打开文档
            abs_path = str(Path(self.file_path).resolve())
            self._word_doc = self._word_app.Documents.Open(abs_path)
            
            # 获取段落并添加批注
            paragraph = self._word_doc.Paragraphs(paragraph_index + 1)
            paragraph.Range.Comments.Add(paragraph.Range, comment_text)
            
            # 保存文档
            self._word_doc.Save()
            
        except Exception as e:
            raise Exception(f"添加批注失败: {str(e)}")
        finally:
            pythoncom.CoUninitialize()
    
    def save_document(self, save_path: str = None):
        """保存文档"""
        if not self.document:
            raise ValueError("文档未加载")
        
        if save_path:
            self.document.save(save_path)
        else:
            self.document.save()
    
    def close(self):
        """关闭文档和Word应用"""
        if self._word_doc:
            try:
                self._word_doc.Close(False)
            except:
                pass
        
        if self._word_app:
            try:
                self._word_app.Quit()
            except:
                pass
        
        self._word_doc = None
        self._word_app = None
        self.document = None
    
    def __enter__(self):
        if self.file_path:
            self.load_document()
        return self
    
    def __exit__(self, exc_type, exc_val, exc_tb):
        self.close()
